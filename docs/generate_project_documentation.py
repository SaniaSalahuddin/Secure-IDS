"""
Generate Word documentation for the Intrusion Detection System project.
Run: python docs/generate_project_documentation.py
"""

import os
from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

OUTPUT = os.path.join(
    os.path.dirname(os.path.abspath(__file__)),
    "Intrusion_Detection_System_Project_Documentation.docx",
)


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    return p


def add_bullet(doc, text):
    doc.add_paragraph(text, style="List Bullet")


def build_document():
    doc = Document()

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("INTRUSION DETECTION SYSTEM\n")
    r.bold = True
    r.font.size = Pt(22)
    title.add_run("Full Project Documentation\n\n").font.size = Pt(14)
    title.add_run(f"Generated: {date.today().strftime('%B %d, %Y')}\n")
    title.add_run("Stack: React + Node.js + FastAPI + MongoDB + Scikit-learn")

    doc.add_page_break()

    sections = [
        ("1. Introduction", """
This project is a web-based Intrusion Detection System (IDS) that uses machine learning to classify network traffic as normal or as a specific attack type. The system is inspired by the NSL-KDD dataset format, which is a standard benchmark for network intrusion detection research.

Users can register and log in with two-factor authentication (OTP via email), upload network traffic data as CSV files or Wireshark captures (PCAP/PCAPNG), and view detection results including attack type, confidence percentage, and per-row predictions. Scan history is stored in MongoDB for later review on the Attack Log page.

The application follows a three-tier architecture: React frontend (port 5173/5174), Node.js Express backend (port 7000), and Python FastAPI ML service (port 8000).
        """),
        ("2. Objectives", """
• Build a secure web portal for intrusion detection analysis.
• Integrate a trained machine learning model (ids_model.pkl) for real-time predictions.
• Support CSV uploads in NSL-KDD feature format.
• Support Wireshark PCAP files with automatic conversion to model-compatible features.
• Implement role-based access (admin, analyst, viewer) for upload permissions.
• Provide OTP-based two-factor authentication for login security.
• Store and display historical attack detection results.
        """),
        ("3. System Architecture", """
The system consists of three main layers:

FRONTEND (React + Vite)
• User interface for login, signup, dashboard, file upload, and attack history.
• Communicates with backend via REST API and JWT tokens stored in localStorage.

BACKEND (Node.js + Express)
• Handles authentication, file uploads, database operations.
• Proxies traffic data to the ML service for prediction.
• Uses MongoDB for users and attack records.

ML SERVICE (Python + FastAPI)
• Loads the trained scikit-learn model (ids_model.pkl).
• Exposes /predict for CSV row data and /predict-pcap for Wireshark files.
• Uses pcap_to_csv.py to convert packet captures into NSL-KDD-style features.

DATA FLOW FOR CSV UPLOAD:
1. User selects CSV → Frontend sends multipart POST to /api/upload
2. Backend reads CSV with csv-parser → sends JSON { data: rows } to http://localhost:8000/predict
3. ML service returns predictions → Backend saves to MongoDB → Frontend displays results

DATA FLOW FOR WIRESHARK UPLOAD:
1. User selects .pcap/.pcapng → Backend forwards file to /predict-pcap
2. ML service converts PCAP to connection rows → runs model → returns results
3. Same response format as CSV for consistent UI display
        """),
        ("4. Technology Stack", """
| Layer      | Technologies |
|------------|--------------|
| Frontend   | React 19, Vite, React Router, Axios, Tailwind CSS, Framer Motion, React Hot Toast |
| Backend    | Node.js, Express 5, MongoDB/Mongoose, Multer, JWT, Bcrypt, Resend (email OTP) |
| ML API     | Python, FastAPI, Uvicorn, Pandas, Scikit-learn, Joblib, Scapy |
| Database   | MongoDB |
| Auth       | JWT + Email OTP (6-digit, 10-minute expiry) |
        """),
        ("5. Project Folder Structure", """
Intrusion-Detection-System/
├── frontend/                 React web application
│   ├── src/
│   │   ├── api/client.js     Axios instance with JWT interceptor
│   │   ├── context/AuthContext.jsx
│   │   ├── components/       Navbar, Sidebar, ProtectedRoute
│   │   ├── layout/DashboardLayout.jsx
│   │   └── pages/            Login, Signup, VerifyOtp, Dashboard, Upload, Attacks
│   └── package.json
├── backend/                  Express API server
│   ├── server.js             Entry point
│   ├── controllers/          authController, uploadController
│   ├── routes/               authRoutes, uploadRoutes, attackRoutes
│   ├── middleware/           authMiddleware, roleMiddleware
│   ├── models/               User.js, Attack.js
│   ├── utils/sendOtp.js      Email via Resend
│   └── uploads/              Temporary uploaded files
├── python/                   ML microservice
│   ├── app.py                FastAPI predict endpoints
│   ├── pcap_to_csv.py        Wireshark → CSV converter
│   ├── ids_model.pkl         Trained ML model
│   └── requirements.txt
├── samples/                  Test CSV and PCAP files
└── docs/                     Documentation
        """),
        ("6. Database Design", """
MongoDB database: Intrusion-Detection (configured in .env as MONGO_URI)

COLLECTION: users
• name (String, required)
• email (String, required, unique)
• password (String, hashed with bcrypt)
• role (enum: admin | analyst | viewer, default: viewer)
• otp (String, temporary for login)
• otpExpiry (Date)

COLLECTION: attacks
• attackType (String) — e.g. normal, neptune, smurf, satan
• confidence (Number) — percentage 0–100
• uploadedBy (ObjectId ref User)
• createdAt (Date, auto)

Only admin and analyst roles can upload files. Viewer can log in but cannot run detection.
        """),
        ("7. Authentication Module", """
SIGNUP (/api/auth/signup)
• Frontend: Signup.jsx collects name, email, password, role.
• Backend: authController.signup hashes password with bcrypt (10 rounds), creates User in MongoDB.

LOGIN (/api/auth/login)
• User submits email + password.
• Backend verifies credentials, generates 6-digit OTP, saves otp + otpExpiry (10 min), sends email via Resend (sendOtp.js).
• User is redirected to Verify OTP page (email stored in localStorage).

VERIFY OTP (/api/auth/verify-otp)
• User enters 6-digit code (VerifyOtp.jsx with react-otp-input).
• Backend validates OTP and expiry, issues JWT token containing userId and role.
• AuthContext stores token and user in localStorage.

PROTECTED ROUTES
• ProtectedRoute.jsx checks AuthContext; redirects to / if not logged in.
• authMiddleware.js reads Authorization: Bearer <token>, verifies with JWT_SECRET.
• roleMiddleware.js restricts /api/upload to admin and analyst only.
        """),
        ("8. Frontend Pages Explained", """
Login.jsx — Email/password form, POST to /api/auth/login, navigates to /verify-otp.

VerifyOtp.jsx — 6-digit OTP input, POST to /api/auth/verify-otp, calls AuthContext.login(), goes to /dashboard.

Signup.jsx — Registration with role selection (viewer default).

Dashboard.jsx — Welcome screen inside DashboardLayout, link to Upload page.

Upload.jsx — Main IDS feature:
• Accepts .csv, .pcap, .pcapng
• Lists sample files from samples/ folder
• Shows detection summary, per-row table, Wireshark badge
• Uses api/client.js for authenticated requests

Attacks.jsx — GET /api/attacks, table of past scans with date, attack type, confidence.

DashboardLayout.jsx — Sidebar + Navbar wrapper for authenticated pages.

Sidebar.jsx — Navigation: Dashboard, Upload Traffic, Attack Log, Manage Users (admin only), Logout.

api/client.js — Base URL http://localhost:7000, attaches Bearer token from localStorage on every request.
        """),
        ("9. Backend API Reference", """
AUTH
POST /api/auth/signup     — Register new user
POST /api/auth/login      — Send OTP to email
POST /api/auth/verify-otp — Verify OTP, return JWT

UPLOAD (requires JWT, role: admin|analyst)
POST /api/upload          — Upload CSV or PCAP file (multipart field: file)

ATTACKS (requires JWT)
GET /api/attacks          — List all attack records (newest first)

ENVIRONMENT (.env)
PORT=7000
MONGO_URI=mongodb://localhost:27017/Intrusion-Detection
JWT_SECRET=your_secret
EMAIL, EMAIL_PASSWORD, RESEND_API (for OTP emails)
ML_API_URL=http://localhost:8000/predict (optional)
        """),
        ("10. Machine Learning Module", """
MODEL FILE: python/ids_model.pkl
• Scikit-learn Pipeline trained on NSL-KDD style features.
• 41 input features (duration, protocol_type, service, flag, byte counts, rates, etc.).
• Output classes include: normal, neptune, smurf, satan, portsweep, nmap, and others.

app.py — FastAPI application
• GET /health — Service health check
• POST /predict — Body: { "data": [ {row1}, {row2}, ... ] }
• POST /predict-pcap — Multipart file upload (.pcap/.pcapng)

run_prediction() function:
1. Builds pandas DataFrame from rows
2. Aligns columns to FEATURE_NAMES (missing columns filled with 0)
3. model.predict() and model.predict_proba() for confidence
4. Summary logic: if ANY row is attack, summary highlights attack type (not majority normal)
5. Returns: success, total_rows, summary, predictions[], time

pcap_to_csv.py — Wireshark converter
• Groups packets into flows (connections) by IP/port
• Extracts: duration, protocol, service (from port), TCP flag, bytes, land, fragments
• Approximates host-level stats (count, dst_host_*, error rates) over entire capture
• CLI: python pcap_to_csv.py capture.pcap -o output.csv
• Function: pcap_path_to_rows() used by /predict-pcap API
        """),
        ("11. NSL-KDD Features (41 Columns)", """
The model expects these columns in each CSV row:

Basic connection: duration, protocol_type (tcp/udp/icmp), service (http, ftp, etc.), flag (SF, REJ, S0...)
Bytes: src_bytes, dst_bytes
Flags: land, wrong_fragment, urgent, hot
Login-related: num_failed_logins, logged_in, is_host_login, is_guest_login
Compromise: num_compromised, root_shell, su_attempted, num_root, num_file_creations, num_shells, num_access_files, num_outbound_cmds
Traffic rates: count, srv_count, serror_rate, srv_serror_rate, rerror_rate, srv_rerror_rate, same_srv_rate, diff_srv_rate, srv_diff_host_rate
Destination host stats: dst_host_count, dst_host_srv_count, dst_host_same_srv_rate, dst_host_diff_srv_rate, dst_host_same_src_port_rate, dst_host_srv_diff_host_rate, dst_host_serror_rate, dst_host_srv_serror_rate, dst_host_rerror_rate, dst_host_srv_rerror_rate

Sample files in samples/ folder are pre-built for testing with known expected labels.
        """),
        ("12. Upload Controller Logic", """
File: backend/controllers/uploadController.js

1. Receives file via Multer (saved to backend/uploads/)
2. If extension is .pcap or .pcapng:
   - Forwards file to ML /predict-pcap using FormData
3. If extension is .csv:
   - Streams file through csv-parser to build rows array
   - POST rows to ML /predict
4. On success:
   - Creates Attack document in MongoDB
   - Returns JSON: attackType, confidence, status, predictions, attacksDetected, safeRows, source
5. Deletes temporary upload file
6. Returns clear error messages from ML validation failures
        """),
        ("13. Sample Test Files", """
| File | Description |
| sample_traffic_test.csv | 2 normal + 1 neptune |
| all_normal.csv | All safe traffic |
| all_attacks.csv | neptune, smurf, satan |
| single_normal.csv | One normal row |
| single_neptune_dos.csv | One DoS attack |
| single_smurf.csv | One smurf attack |
| mixed_5_rows.csv | Mixed normal and attacks |
| dos_and_probe.csv | Multiple neptune + satan |
| test_capture.pcap | Small Wireshark test capture |

Use these from Upload Traffic page to demonstrate the system in reports and viva.
        """),
        ("14. How to Run the Project", """
PREREQUISITES
• Node.js 18+
• Python 3.10+
• MongoDB running locally

STEP 1 — MongoDB
Start MongoDB service (default port 27017).

STEP 2 — Python ML API
cd python
pip install -r requirements.txt
uvicorn app:app --reload --port 8000

STEP 3 — Backend
cd backend
npm install
npm start
(Server runs on port 7000)

STEP 4 — Frontend
cd frontend
npm install
npm run dev
(Open http://localhost:5173)

STEP 5 — Create user
• Sign up at /signup (choose role: analyst or admin for upload access)
• Log in, enter OTP from email
• Go to Upload Traffic and test with samples/*.csv
        """),
        ("15. Security Features", """
• Passwords hashed with bcrypt before storage
• JWT tokens expire after 1 day
• OTP expires after 10 minutes
• Role-based upload restriction (admin, analyst only)
• Protected API routes require valid Bearer token
• CORS enabled for frontend-backend communication
• Uploaded files deleted after processing from server disk

Note: For production, use HTTPS, secure JWT secrets, rate limiting, and environment variable protection.
        """),
        ("16. Limitations and Future Work", """
LIMITATIONS
• Wireshark PCAP conversion approximates NSL-KDD features; predictions on real captures are best-effort.
• Model trained on NSL-KDD; may not generalize to all modern attack types.
• OTP email requires Resend API configuration.
• Manage Users page linked in sidebar but route not fully implemented.

FUTURE ENHANCEMENTS
• Real-time packet capture integration
• Full KDD 2-second window feature extraction from PCAP
• Charts and analytics on Attack Log page
• Export reports as PDF
• Admin user management UI
• Model retraining pipeline from new labeled data
        """),
        ("17. Conclusion", """
The Intrusion Detection System successfully combines a modern web stack with machine learning to detect network attacks from uploaded traffic data. The three-service architecture (React, Express, FastAPI) keeps concerns separated and makes the system maintainable and scalable.

Key achievements include secure OTP login, CSV and Wireshark upload support, per-row and summary detection display, and MongoDB persistence of scan history. The project demonstrates practical application of the NSL-KDD dataset, scikit-learn deployment, and full-stack web development suitable for academic submission and demonstration.
        """),
    ]

    for heading, body in sections:
        add_heading(doc, heading, level=1)
        for paragraph in body.strip().split("\n\n"):
            paragraph = paragraph.strip()
            if paragraph.startswith("|") or paragraph.startswith("•"):
                for line in paragraph.split("\n"):
                    line = line.strip()
                    if line.startswith("•"):
                        add_bullet(doc, line[1:].strip())
                    elif line.startswith("|") and "---" not in line:
                        add_para(doc, line)
            else:
                add_para(doc, paragraph)

    doc.add_page_break()
    add_heading(doc, "Appendix A: Key File Paths Quick Reference", level=1)
    paths = [
        "frontend/src/pages/Upload.jsx — Upload UI",
        "frontend/src/api/client.js — API client",
        "backend/server.js — Server entry",
        "backend/controllers/uploadController.js — File upload logic",
        "backend/controllers/authController.js — Auth logic",
        "python/app.py — ML API",
        "python/pcap_to_csv.py — PCAP converter",
        "python/ids_model.pkl — Trained model",
        "samples/ — Test data files",
    ]
    for p in paths:
        add_bullet(doc, p)

    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    path = build_document()
    print(f"Created: {path}")
